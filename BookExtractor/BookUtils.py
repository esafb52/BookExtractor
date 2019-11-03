import re
import os
from docx import Document
from docx.shared import Inches
from docx.shared import Pt


class BookUtils:

    @staticmethod
    def remove_all_none_aplefba_char(my_str):
        txt = str(my_str)
        final = re.sub(r"[^\w ]+", '', txt)
        return final

    @staticmethod
    def remove_all_english_aplefba(my_str):
        txt = str(my_str)
        final = re.sub(r"[a-zA-Z]+", '', txt)
        return final

    @staticmethod
    def remove_all_digit(my_str):
        txt = str(my_str)
        final = re.sub(r"[0-9]+", '', txt)
        return final

    @staticmethod
    def remove_w_space(txt):
        temp_txt = str(txt)
        res = re.sub('[  ]+', ' ', temp_txt)
        return res

    @staticmethod
    def make_dir(out_dir):
        if not os.path.exists(out_dir):
            os.makedirs(out_dir)

    @staticmethod
    def convert_docx_book_one_head_many_para(out_dir, lst_res, out_file_name, cover_pic):
        BookUtils.make_dir(out_dir)
        my_book = Document()
        core_properties = my_book.core_properties
        core_properties.author = 'esa shahgolizadeh'
        style = my_book.styles['Normal']
        font = style.font
        font.name = 'B Lotus'
        font.size = Pt(14)
        if os.path.isfile(cover_pic):
            my_book.add_picture(cover_pic, width=Inches(6.25))
        try:
            for row in lst_res:
                my_book.add_heading(row['head'], level=1)
                all_para = row['all_txt']
                for para in all_para:
                    paragraph = my_book.add_paragraph(para)
                    paragraph.alignment = 1
            out_file = os.path.join(out_dir, out_file_name)
            print(out_file)
            my_book.save(out_file)
        except Exception as e:
            print('error!!!', e)

    @staticmethod
    def convert_docx_book_one_head_many_para(out_dir, lst_res, out_file_name):
        BookUtils.make_dir(out_dir)
        my_book = Document()
        core_properties = my_book.core_properties
        core_properties.author = 'esa shahgolizadeh'
        style = my_book.styles['Normal']
        font = style.font
        font.name = 'B Lotus'
        font.size = Pt(14)
        try:
            if len(lst_res) <= 0:
                print('no item found!!!')
                return
            for row in lst_res:
                my_book.add_heading(row['head'], level=1)
                all_para = row['all_txt']
                for para in all_para:
                    my_book.add_paragraph(para)
            out_file = os.path.join(out_dir, out_file_name)
            print(out_file)
            my_book.save(out_file)
        except Exception as e:
            print('error!!!', e)


@staticmethod
def convert_docx_book(out_dir, lst_res, out_file_name):
    if not os.path.exists(out_dir):
        os.makedirs(out_dir)
    my_book = Document()
    try:
        if len(lst_res) <= 0:
            print('no item found!!!')
            return
        for row in lst_res:
            if row['type'] == 'head':
                my_book.add_heading(row['txt'], level=1)
            if row['type'] == 'pic':
                my_book.add_picture(row['pic'], width=Inches(6.25))
            if row['type'] == 'para':
                my_book.add_paragraph(row['txt'])
        out_file = os.path.join(out_dir, out_file_name)
        print(out_file)
        my_book.save(out_file)
    except Exception as e:
        print(e)
